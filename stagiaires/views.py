from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from django.http import HttpResponse
from .models import Stage, Application, Stagiaire, ProgressUpdate, Projet_realise, Stagiaire, Encadreur, FileUpload, Absence, CertificateRequest
from .forms import ProgressUpdateForm, ProjetRealiseForm, FileUploadForm, StagiaireProfileForm, CertificateRequestForm
import docx
import tempfile
import pypandoc


@login_required
def stagiaire_dashboard(request):
    if request.user.user_type != 1:
        return redirect('login')

    stagiaire = Stagiaire.objects.get(email=request.user.email)  # email is unique
    stages_disponibles = Stage.objects.all()

    for stage in stages_disponibles:
        stage.already_applied = Application.objects.filter(stagiaire=stagiaire, stage=stage).exists()

    mes_candidatures = Application.objects.filter(stagiaire=stagiaire)
    mes_stages_actuels = Stage.objects.filter(application__stagiaire=stagiaire, application__statut='accepté')
    progress_updates = ProgressUpdate.objects.filter(stagiaire=stagiaire)
    projets_realises = Projet_realise.objects.filter(stagiaire=stagiaire)
    rapport_files = FileUpload.objects.filter(user=request.user, file_type=FileUpload.TRAINEE_REPORT)
    presentation_files = FileUpload.objects.filter(user=request.user, file_type=FileUpload.PRESENTATION)
    convention_files = FileUpload.objects.filter(user=request.user, file_type=FileUpload.STAGE_CONVENTION)


    encadreur = None
    accepted_application = Application.objects.filter(stagiaire=stagiaire, statut='accepté').first()
    if accepted_application:
        encadreur = accepted_application.encadreur


    absences_count = Absence.objects.filter(stagiaire=stagiaire, status='absent').count()
    retards_count = Absence.objects.filter(stagiaire=stagiaire, status='retard').count()

    context = {
        'stages_disponibles': stages_disponibles,
        'mes_candidatures': mes_candidatures,
        'mes_stages_actuels': mes_stages_actuels,
        'progress_updates': progress_updates,
        'projets_realises': projets_realises,
        'rapport_files_count': rapport_files.count(),
        'presentation_files_count': presentation_files.count(),
        'convention_files_count': convention_files.count(),
        'encadreur': encadreur,
        'absences_count': absences_count,
        'retards_count': retards_count,
    }

    return render(request, 'stagiaires/dashboard.html', context)


@login_required
def postuler_stage(request, stage_id):
    if request.user.user_type != 1:
        return redirect('login')
    stage = get_object_or_404(Stage, id=stage_id)
    stagiaire = Stagiaire.objects.get(email=request.user.email)  # email is unique
    

    if not Application.objects.filter(stagiaire=stagiaire, stage=stage).exists():
        Application.objects.create(stagiaire=stagiaire, stage=stage)
    
    return redirect('stagiaire_dashboard')

@login_required
def details_stage(request, stage_id):
    stage = get_object_or_404(Stage, id=stage_id)
    return render(request, 'stagiaires/details_stage.html', {'stage': stage})

@login_required
def mes_candidatures(request):
    if request.user.user_type != 1:
        return redirect('login')
    stagiaire = Stagiaire.objects.get(email=request.user.email)
    mes_candidatures = Application.objects.filter(stagiaire=stagiaire)
    return render(request, 'stagiaires/mes_candidatures.html', {
        'mes_candidatures': mes_candidatures,
    })

@login_required
def delete_application(request, application_id):
    if request.user.user_type != 1:
        return redirect('login')
    
    application = get_object_or_404(Application, id=application_id, stagiaire__email=request.user.email)
    application.delete()
    return redirect('mes_candidatures')


@login_required
def stages_disponibles(request):
    if request.user.user_type != 1:
        return redirect('login')

    stagiaire = Stagiaire.objects.get(email=request.user.email)
    stages_disponibles = Stage.objects.all()

    for stage in stages_disponibles:
        stage.already_applied = Application.objects.filter(stagiaire=stagiaire, stage=stage).exists()

    return render(request, 'stagiaires/stages_disponibles.html', {
        'stages_disponibles': stages_disponibles
    })

@login_required
def suivi_de_stage(request):
    if request.user.user_type != 1:
        return redirect('login')

    stagiaire = Stagiaire.objects.get(email=request.user.email)
    progress_updates = ProgressUpdate.objects.filter(stagiaire=stagiaire)

    if request.method == 'POST':
        form = ProgressUpdateForm(request.POST)
        if form.is_valid():
            progress_update = form.save(commit=False)
            progress_update.stagiaire = stagiaire
            progress_update.save()
            return redirect('suivi_de_stage')
    else:
        form = ProgressUpdateForm()

    return render(request, 'stagiaires/suivi_de_stage.html', {
        'progress_updates': progress_updates,
        'form': form
    })

@login_required
def soumettre_projet(request):
    if request.user.user_type != 1:
        return redirect('login')

    stagiaire = Stagiaire.objects.get(email=request.user.email)
    encadreur = Encadreur.objects.first()

    if request.method == 'POST':
        form = ProjetRealiseForm(request.POST, request.FILES)
        if form.is_valid():
            projet = form.save(commit=False)
            projet.stagiaire = stagiaire
            projet.encadreur = encadreur
            projet.save()
            return redirect('stagiaire_dashboard')
    else:
        form = ProjetRealiseForm()

    return render(request, 'stagiaires/soumettre_projet.html', {'form': form})

@login_required
def upload_files(request):
    if request.method == 'POST':
        form = FileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            upload = form.save(commit=False)
            upload.user = request.user
            if upload.is_valid_extension():
                upload.save()
                return redirect('upload_files')
            else:
                form.add_error('file', 'Invalid file extension.')
    else:
        form = FileUploadForm()
    

    rapport_files = FileUpload.objects.filter(user=request.user, file_type=FileUpload.TRAINEE_REPORT)
    presentation_files = FileUpload.objects.filter(user=request.user, file_type=FileUpload.PRESENTATION)
    convention_files = FileUpload.objects.filter(user=request.user, file_type=FileUpload.STAGE_CONVENTION)

    context = {
        'form': form,
        'rapport_files': rapport_files,
        'presentation_files': presentation_files,
        'convention_files': convention_files,
    }
    return render(request, 'stagiaires/upload_files.html', context)

def profile(request):
    stagiaire = Stagiaire.objects.get(user=request.user)
    
    if request.method == 'POST':
        form = StagiaireProfileForm(request.POST, instance=stagiaire)
        if form.is_valid():
            form.save()
            return redirect('profile')
    else:
        form = StagiaireProfileForm(instance=stagiaire)

    return render(request, 'stagiaires/profile.html', {'form': form})

@login_required
def request_certificate(request):
    certificate_request = None
    if request.method == 'POST':
        form = CertificateRequestForm(request.POST)
        if form.is_valid():
            certificate_request = form.save(commit=False)
            certificate_request.stagiaire = request.user.stagiaire
            certificate_request.save()
            return redirect('request_certificate')
    else:
        form = CertificateRequestForm()
        certificate_request = CertificateRequest.objects.filter(stagiaire=request.user.stagiaire).last()

    return render(request, 'stagiaires/request_certificate.html', {
        'form': form,
        'certificate_request': certificate_request
    })
@login_required
def download_certificate(request, request_id):
    certificate_request = get_object_or_404(CertificateRequest, id=request_id, stagiaire=request.user.stagiaire, status=CertificateRequest.APPROVED)

    # Generate the Word document
    doc = docx.Document()
    doc.add_heading('Certificat de Stage', 0)
    doc.add_paragraph(f"Ceci est pour certifier que {request.user.stagiaire.nom} {request.user.stagiaire.prenom} a complété avec succès le stage.")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_doc:
        doc_path = tmp_doc.name
        doc.save(doc_path)


    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
        pdf_path = tmp_pdf.name


    try:
        pypandoc.convert_file(doc_path, 'pdf', outputfile=pdf_path, extra_args=['--pdf-engine=xelatex'])
    except OSError:
        pypandoc.download_pandoc()
        pypandoc.convert_file(doc_path, 'pdf', outputfile=pdf_path, extra_args=['--pdf-engine=xelatex'])

    with open(pdf_path, 'rb') as pdf:
        response = HttpResponse(pdf.read(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename=certificate_{request_id}.pdf'
        return response